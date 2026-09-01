import os
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment, PatternFill
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime, timedelta
from database.db import (
    get_all_users_raw, get_all_orders_raw, get_all_transactions_raw, 
    get_global_stats, get_orders_by_date_range, get_users_by_role_raw,
    get_user_full_report, get_all_starters_with_dates
)

async def generate_starters_report_word(output_path):
    """Generates a Word document with all user IDs and their start dates."""
    from database.db import get_all_starters_with_dates
    starters = await get_all_starters_with_dates()
    
    doc = Document()
    title = doc.add_heading('CHIROQCHI TAKSI - FOYDALANUVCHILAR RO\'YXATI', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p = doc.add_paragraph()
    p.add_run(f"Jami foydalanuvchilar: ").bold = True
    p.add_run(f"{len(starters)} ta\n")
    p.add_run(f"Eksport vaqti: {datetime.now().strftime('%d.%m.%Y %H:%M')}")
    
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '№'
    hdr_cells[1].text = 'Telegram ID'
    hdr_cells[2].text = 'Start vaqti'
    
    for i, (uid, ts) in enumerate(starters, 1):
        row_cells = table.add_row().cells
        row_cells[0].text = str(i)
        row_cells[1].text = str(uid)
        row_cells[2].text = format_db_date(str(ts)) if ts else "Noma'lum"
        
    doc.save(output_path)
    return output_path

def format_db_date(date_str):
    """Helper to format SQLite timestamp strings to human readable format."""
    try:
        if not date_str:
            return "Noma'lum"
        dt = datetime.strptime(date_str, "%Y-%m-%d %H:%M:%S")
        return dt.strftime("%d.%m.%Y %H:%M")
    except:
        return str(date_str)

def apply_header_style(ws):
    """Apply professional header styling to worksheet."""
    header_fill = PatternFill(start_color="1a73e8", end_color="1a73e8", fill_type="solid")
    for cell in ws[1]:
        cell.font = Font(bold=True, color="FFFFFF", size=11)
        cell.fill = header_fill
        cell.alignment = Alignment(horizontal="center", vertical="center")
    ws.row_dimensions[1].height = 20

def auto_adjust_columns(ws):
    """Auto-adjust all column widths."""
    for col in ws.columns:
        max_length = 0
        column = col[0].column_letter
        for cell in col:
            try:
                if len(str(cell.value)) > max_length:
                    max_length = len(str(cell.value))
            except:
                pass
        ws.column_dimensions[column].width = min(max_length + 3, 40)

def parse_date_input(text: str):
    """
    Parse admin date input. Supports:
    - 'today' / 'bugun'
    - 'YYYY-MM-DD' (daily)
    - 'YYYY-MM' (monthly)
    - 'YYYY' (yearly)
    Returns (start_date, end_date, label) as YYYY-MM-DD strings.
    """
    text = text.strip().lower()
    today = datetime.now()
    
    if text in ('today', 'bugun'):
        d = today.strftime('%Y-%m-%d')
        return d, d, f"Bugun ({d})"
    
    # YYYY-MM-DD: exact day
    if len(text) == 10 and text.count('-') == 2:
        try:
            datetime.strptime(text, '%Y-%m-%d')
            return text, text, f"Kun: {text}"
        except:
            pass
    
    # YYYY-MM: full month
    if len(text) == 7 and text.count('-') == 1:
        try:
            dt = datetime.strptime(text, '%Y-%m')
            # Last day of month
            if dt.month == 12:
                end_dt = dt.replace(day=31)
            else:
                end_dt = dt.replace(month=dt.month+1, day=1) - timedelta(days=1)
            start = dt.strftime('%Y-%m-%d')
            end = end_dt.strftime('%Y-%m-%d')
            return start, end, f"Oy: {dt.strftime('%B %Y')}"
        except:
            pass
    
    # YYYY: full year
    if len(text) == 4 and text.isdigit():
        year = int(text)
        return f"{year}-01-01", f"{year}-12-31", f"Yil: {year}"
    
    return None, None, None

# --- DATE-FILTERED REPORTS ---

async def generate_date_report_excel(output_path, start_date, end_date, label):
    """Generate Excel report for orders and users in a date range."""
    wb = Workbook()
    header_fill = PatternFill(start_color="1a73e8", end_color="1a73e8", fill_type="solid")

    # === ORDERS SHEET ===
    ws_orders = wb.active
    ws_orders.title = "Buyurtmalar"
    order_headers = ["№", "Buyurtma ID", "Yo'lovchi ID", "Haydovchi ID", "Qayerdan", "Qayerga", "Narx", "Holat", "Tur", "Vaqt"]
    ws_orders.append(order_headers)
    apply_header_style(ws_orders)
    
    orders = await get_orders_by_date_range(start_date, end_date)
    orders.sort(key=lambda x: x[16] if x[16] else "")
    for i, o in enumerate(orders, 1):
        ws_orders.append([
            i,
            o[0],   # order_id
            o[1],   # passenger_id
            o[2] if o[2] else "-",   # driver_id
            o[3],   # from_location
            o[4],   # to_location
            f"{o[5]:,}" if o[5] else "0",  # price
            o[11],  # status (Corrected index)
            o[13] if len(o) > 13 else "taxi",  # order_type (Corrected index)
            format_db_date(str(o[16]))  # timestamp (Corrected index)
        ])
    auto_adjust_columns(ws_orders)
    ws_orders.append([])
    ws_orders.append([f"Jami buyurtmalar: {len(orders)} ta | Davr: {label}"])

    # === USERS SHEET ===
    ws_users = wb.create_sheet(title="Yangi Foydalanuvchilar")
    user_headers = ["№", "ID", "F.I.SH", "Telefon", "Rol", "Holat", "Balans", "Ro'yxat Sanasi"]
    ws_users.append(user_headers)
    apply_header_style(ws_users)
    
    all_users = await get_all_users_raw()
    new_users = [u for u in all_users if u[10] and start_date <= u[10][:10] <= end_date]
    for i, u in enumerate(new_users, 1):
        ws_users.append([
            i, u[0], u[1], u[2],
            "Haydovchi" if u[4] == 'driver' else "Yo'lovchi",
            u[5], f"{u[6]:,}" if u[6] else "0",
            format_db_date(str(u[10]))
        ])
    auto_adjust_columns(ws_users)
    ws_users.append([])
    ws_users.append([f"Yangi foydalanuvchilar: {len(new_users)} ta | Davr: {label}"])

    # === SUMMARY SHEET ===
    ws_sum = wb.create_sheet(title="Xulosa")
    ws_sum.append(["CHIROQCHI TAKSI — HISOBOT"])
    ws_sum["A1"].font = Font(bold=True, size=14, color="1a73e8")
    ws_sum.append(["Davr:", label])
    ws_sum.append(["Hisobot vaqti:", datetime.now().strftime("%d.%m.%Y %H:%M")])
    ws_sum.append([])
    ws_sum.append(["Ko'rsatkich", "Qiymat"])
    ws_sum.append(["Jami buyurtmalar", len(orders)])
    finished = [o for o in orders if o[11] == 'finished']
    cancelled = [o for o in orders if o[11] == 'cancelled']
    total_rev = sum(o[5] for o in finished if o[5])
    ws_sum.append(["Yakunlangan buyurtmalar", len(finished)])
    ws_sum.append(["Bekor qilingan buyurtmalar", len(cancelled)])
    ws_sum.append(["Umumiy tushum (so'm)", f"{total_rev:,}"])
    ws_sum.append(["Yangi foydalanuvchilar", len(new_users)])
    auto_adjust_columns(ws_sum)

    wb.save(output_path)
    return output_path

# --- ROLE-FILTERED REPORTS ---

async def generate_role_report_excel(output_path, role):
    """Generate Excel report for all users of a specific role."""
    wb = Workbook()
    label = "Haydovchilar" if role == 'driver' else "Yo'lovchilar"
    
    ws = wb.active
    ws.title = label
    
    if role == 'driver':
        headers = ["№", "ID", "F.I.SH", "Telefon", "Holat", "Balans", "Mashina", "Raqam", "Reyting", "Safarlar", "Ro'yxat Sanasi"]
    else:
        headers = ["№", "ID", "F.I.SH", "Telefon", "Holat", "Balans", "Jami Buyrutmalar", "Ro'yxat Sanasi"]
    
    ws.append(headers)
    apply_header_style(ws)
    
    users = await get_users_by_role_raw(role)
    all_orders = await get_all_orders_raw()
    
    for i, u in enumerate(users, 1):
        if role == 'driver':
            ws.append([
                i, u[0], u[1], u[2], u[5],
                f"{u[6]:,}" if u[6] else "0",
                u[11] if u[11] else "-",  # car_name
                u[12] if u[12] else "-",  # car_number
                f"{u[14]:.1f}" if u[14] else "0.0",  # rating
                u[15] if u[15] else 0, # total_rides
                format_db_date(str(u[10]))
            ])
        else:
            order_count = sum(1 for o in all_orders if o[1] == u[0])
            ws.append([
                i, u[0], u[1], u[2], u[5],
                f"{u[6]:,}" if u[6] else "0",
                order_count,
                format_db_date(str(u[10]))
            ])
    
    auto_adjust_columns(ws)
    ws.append([])
    ws.append([f"Jami {label}: {len(users)} ta"])
    
    wb.save(output_path)
    return output_path

# --- USER SPECIFIC REPORT ---

async def generate_user_report_excel(output_path, user_id: int):
    """Generate a detailed report for a specific user."""
    user, orders, transactions, driver_info = await get_user_full_report(user_id)
    
    if not user:
        return None
    
    wb = Workbook()
    role_label = "Haydovchi" if user[4] == 'driver' else "Yo'lovchi"
    
    # === PROFILE SHEET ===
    ws_profile = wb.active
    ws_profile.title = "Profil"
    ws_profile.append(["CHIROQCHI TAKSI — FOYDALANUVCHI HISOBOTI"])
    ws_profile["A1"].font = Font(bold=True, size=14, color="1a73e8")
    ws_profile.append([])
    ws_profile.append(["Maydon", "Qiymat"])
    ws_profile.append(["User ID", user[0]])
    ws_profile.append(["F.I.SH", user[1]])
    ws_profile.append(["Telefon", user[2]])
    ws_profile.append(["Rol", role_label])
    ws_profile.append(["Holat", user[5]])
    ws_profile.append(["Balans", f"{user[6]:,} so'm" if user[6] else "0 so'm"])
    ws_profile.append(["Ro'yxatdan o'tgan", format_db_date(str(user[10]))])
    ws_profile.append(["Jami buyurtmalar", len(orders)])
    ws_profile.append(["Jami tranzaksiyalar", len(transactions)])
    
    if driver_info:
        ws_profile.append([])
        ws_profile.append(["=== HAYDOVCHI MA'LUMOTLARI ==="])
        ws_profile.append(["Mashina nomi", driver_info[1] if driver_info else "-"])
        ws_profile.append(["Mashina raqami", driver_info[2] if driver_info else "-"])
        ws_profile.append(["Reyting", f"{driver_info[4]:.1f}" if driver_info and driver_info[4] else "0.0"])
        ws_profile.append(["Jami safarlar", driver_info[5] if driver_info and len(driver_info) > 5 else 0])
    
    auto_adjust_columns(ws_profile)
    
    # === ORDERS SHEET ===
    ws_orders = wb.create_sheet(title="Buyurtmalar")
    order_headers = ["№", "ID", "Rol", "Qayerdan", "Qayerga", "Narx", "Holat", "Vaqt"]
    ws_orders.append(order_headers)
    apply_header_style(ws_orders)
    
    for i, o in enumerate(orders, 1):
        role_in_order = "Yo'lovchi" if o[1] == user_id else "Haydovchi"
        ws_orders.append([
            i, o[0], role_in_order,
            o[3], o[4],
            f"{o[5]:,}" if o[5] else "0",
            o[11],
            format_db_date(str(o[16]))
        ])
    auto_adjust_columns(ws_orders)
    
    # === TRANSACTIONS SHEET ===
    ws_trans = wb.create_sheet(title="Tranzaksiyalar")
    trans_headers = ["№", "ID", "Summa", "Tur", "Tavsif", "Vaqt"]
    ws_trans.append(trans_headers)
    apply_header_style(ws_trans)
    
    for i, t in enumerate(transactions, 1):
        ws_trans.append([
            i, t[0],
            f"{t[2]:,}" if t[2] else "0",
            "Kirim ✅" if t[4] == 'in' else "Chiqim ❌",
            t[3],
            format_db_date(str(t[5]))
        ])
    auto_adjust_columns(ws_trans)
    
    wb.save(output_path)
    return output_path

# --- ORIGINAL FULL REPORTS (kept for backward compatibility) ---

async def generate_excel_report(output_path):
    wb = Workbook()
    header_fill = PatternFill(start_color="CCE5FF", end_color="CCE5FF", fill_type="solid")
    
    ws_users = wb.active
    ws_users.title = "Foydalanuvchilar"
    headers = ["ID", "F.I.SH", "Tel", "Rol", "Holat", "Balans", "Reg Sana"]
    ws_users.append(headers)
    for cell in ws_users[1]:
        cell.font = Font(bold=True)
        cell.fill = header_fill
        cell.alignment = Alignment(horizontal="center")
    users = await get_all_users_raw()
    for u in users:
        u_list = [u[0], u[1], u[2], u[4], u[5], u[6], u[10]]
        u_list[6] = format_db_date(str(u_list[6])) if u_list[6] else "-"
        ws_users.append(u_list)
        
    ws_orders = wb.create_sheet(title="Buyurtmalar")
    order_headers = ["ID", "Yo'lovchi", "Haydovchi", "Qayerdan", "Qayerga", "Narx", "Holat", "Sana"]
    ws_orders.append(order_headers)
    for cell in ws_orders[1]:
        cell.font = Font(bold=True)
        cell.fill = header_fill
    orders = await get_all_orders_raw()
    for o in orders:
        o_list = [o[0], o[1], o[2], o[3], o[4], o[5], o[11], o[16]]
        o_list[7] = format_db_date(str(o_list[7])) if o_list[7] else "-"
        ws_orders.append(o_list)

    ws_trans = wb.create_sheet(title="Tranzaksiyalar")
    trans_headers = ["ID", "User", "Summa", "Tavsif", "Tur", "Sana"]
    ws_trans.append(trans_headers)
    for cell in ws_trans[1]:
        cell.font = Font(bold=True)
        cell.fill = header_fill
    trans = await get_all_transactions_raw()
    for t in trans:
        t_list = list(t[:6])
        t_list[5] = format_db_date(str(t_list[5])) if t_list[5] else "-"
        ws_trans.append(t_list)

    for ws in wb.worksheets:
        for col in ws.columns:
            max_length = 0
            column = col[0].column_letter
            for cell in col:
                try:
                    if len(str(cell.value)) > max_length:
                        max_length = len(str(cell.value))
                except:
                    pass
            ws.column_dimensions[column].width = max_length + 2

    wb.save(output_path)
    return output_path

async def generate_word_report(output_path):
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)
    
    title = doc.add_heading('CHIROQCHI TAKSI - HISOBOT', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_heading('Umumiy Statistika', level=1)
    stats = await get_global_stats()
    
    p = doc.add_paragraph()
    p.add_run(f"Hisobot vaqti: ").bold = True
    p.add_run(f"{datetime.now().strftime('%d.%m.%Y %H:%M')}\n")
    p.add_run(f"Foydalanuvchilar: ").bold = True
    p.add_run(f"{stats['total_users']} ta\n")
    p.add_run(f"Faol haydovchilar: ").bold = True
    p.add_run(f"{stats['active_drivers']} ta\n")
    p.add_run(f"Jami buyurtmalar: ").bold = True
    p.add_run(f"{stats['total_orders']} ta\n")
    p.add_run(f"Umumiy tushum: ").bold = True
    p.add_run(f"{stats['total_revenue']:,} so'm\n")
    
    doc.add_heading("So'nggi buyurtmalar (50 ta)", level=1)
    orders = await get_all_orders_raw()
    sorted_orders = sorted(orders, key=lambda x: str(x[16]), reverse=True)[:50]
    
    table_o = doc.add_table(rows=1, cols=5)
    table_o.style = 'Table Grid'
    hdr_o = table_o.rows[0].cells
    hdr_o[0].text = "Yo'lovchi"
    hdr_o[1].text = 'Haydovchi'
    hdr_o[2].text = 'Narx'
    hdr_o[3].text = 'Holat'
    hdr_o[4].text = 'Vaqti'
    
    for o in sorted_orders:
        row = table_o.add_row().cells
        row[0].text = str(o[1])
        row[1].text = str(o[2]) if o[2] else "-"
        row[2].text = f"{o[5]:,} s" if o[5] else "0"
        row[3].text = str(o[11])
        row[4].text = format_db_date(str(o[16]))

    doc.save(output_path)
    return output_path
